from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt

from common import active_profile, choose_output_path, read_json, unique_output_path
from inspect_docx import classify_paragraphs, inspect_document
from privacy_scrub import scrub_docx
from validate_docx import validate_document


ROLE_NAMES = {
    "main_title": "ODF Main Title",
    "heading1": "ODF Heading 1",
    "heading2": "ODF Heading 2",
    "body": "ODF Body",
    "reference_note": "ODF Reference Note",
    "description": "ODF Description",
}
ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
H1_RE = re.compile(r"^[一二三四五六七八九十百]+、")
H2_RE = re.compile(r"^（[一二三四五六七八九十百]+）")
NOTE_RE = re.compile(r"^(?:（.*）|\(.*\))$", re.DOTALL)


def installed_font_names() -> set[str]:
    names: set[str] = set()
    if sys.platform != "win32":
        return names
    try:
        import winreg
    except ImportError:
        return names
    locations = [
        (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts"),
        (winreg.HKEY_CURRENT_USER, r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts"),
    ]
    for hive, key_name in locations:
        try:
            with winreg.OpenKey(hive, key_name) as key:
                index = 0
                while True:
                    try:
                        value_name, value, _ = winreg.EnumValue(key, index)
                    except OSError:
                        break
                    names.add(re.sub(r"\s*\([^)]*\)\s*$", "", value_name).strip())
                    names.add(Path(str(value)).stem)
                    index += 1
        except OSError:
            continue
    return names


def normalize_font(value: str) -> str:
    return re.sub(r"[\s_-]+", "", value).casefold()


def resolve_font(preferred: str, fallback: str, installed: set[str], warnings: list[str]) -> str:
    normalized = {normalize_font(name): name for name in installed}
    aliases = {
        normalize_font("黑体"): ["SimHei"],
        normalize_font("楷体_GB2312"): ["楷体GB2312", "KaiTi_GB2312", "KaiTi"],
        normalize_font("仿宋_GB2312"): ["仿宋GB2312", "FangSong_GB2312", "FangSong"],
        normalize_font("宋体"): ["SimSun"],
    }
    for candidate in [preferred, fallback, *aliases.get(normalize_font(fallback), [])]:
        if normalize_font(candidate) in normalized:
            if candidate != preferred:
                warnings.append(f"Font fallback used: {preferred} -> {candidate}")
            return candidate
    warnings.append(f"Font unavailable for visual verification: {preferred}; the preferred name was retained")
    return preferred


def set_rfonts(run_or_font, east_asia: str, latin: str) -> None:
    element = getattr(run_or_font, "_element", None)
    if element is None:
        element = getattr(run_or_font, "_parent", None)
    if hasattr(run_or_font, "_element") and run_or_font.__class__.__name__ == "Run":
        r_pr = run_or_font._element.get_or_add_rPr()
    else:
        r_pr = run_or_font._element.get_or_add_rPr() if hasattr(run_or_font, "_element") else None
    if r_pr is None:
        return
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:cs"), latin)


def set_paragraph_geometry(paragraph, spec: dict) -> None:
    paragraph.alignment = ALIGNMENTS[spec["alignment"]]
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(spec["space_before_pt"])
    fmt.space_after = Pt(0)
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.get_or_add_spacing()
    spacing.set(qn("w:line"), str(round(spec["line_spacing_pt"] * 20)))
    spacing.set(qn("w:lineRule"), "exact")
    spacing.set(qn("w:before"), str(round(spec["space_before_pt"] * 20)))
    spacing.set(qn("w:after"), "0")
    ind = p_pr.get_or_add_ind()
    for attr in ("firstLine", "hanging"):
        key = qn(f"w:{attr}")
        if key in ind.attrib:
            del ind.attrib[key]
    ind.set(qn("w:firstLineChars"), str(round(spec["first_line_chars"] * 100)))


def ensure_styles(document: Document, profile: dict, fonts: dict[str, str]) -> None:
    for role, style_name in ROLE_NAMES.items():
        spec = profile["styles"][role]
        try:
            style = document.styles[style_name]
        except KeyError:
            style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = spec["font_latin"]
        style.font.size = Pt(spec["size_pt"])
        style.font.bold = profile["global"]["bold"]
        set_rfonts(style, fonts[role], spec["font_latin"])
        style.paragraph_format.alignment = ALIGNMENTS[spec["alignment"]]
        style.paragraph_format.space_before = Pt(spec["space_before_pt"])
        style.paragraph_format.space_after = Pt(0)


def apply_role(paragraph, role: str, profile: dict, fonts: dict[str, str]) -> None:
    if role == "skip":
        return
    if role not in ROLE_NAMES:
        raise ValueError(f"Unsupported paragraph role: {role}")
    spec = profile["styles"][role]
    paragraph.style = ROLE_NAMES[role]
    set_paragraph_geometry(paragraph, spec)
    for run in paragraph.runs:
        run.font.name = spec["font_latin"]
        run.font.size = Pt(spec["size_pt"])
        run.font.bold = profile["global"]["bold"]
        set_rfonts(run, fonts[role], spec["font_latin"])


def apply_page_setup(document: Document, profile: dict, warnings: list[str]) -> None:
    margins = profile["page"]["margins_cm"]
    for index, section in enumerate(document.sections):
        if section.orientation == WD_ORIENT.LANDSCAPE:
            warnings.append(f"Landscape section {index + 1} was preserved")
            continue
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Cm(margins["top"])
        section.bottom_margin = Cm(margins["bottom"])
        section.left_margin = Cm(margins["left"])
        section.right_margin = Cm(margins["right"])


def clear_paragraph(paragraph) -> None:
    p_pr = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not p_pr:
            paragraph._p.remove(child)


def add_page_field(paragraph, spec: dict, font_name: str) -> None:
    clear_paragraph(paragraph)
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    left = paragraph.add_run("— ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    field_run = OxmlElement("w:r")
    field_text = OxmlElement("w:t")
    field_text.text = "1"
    field_run.append(field_text)
    field.append(field_run)
    paragraph._p.append(field)
    right = paragraph.add_run(" —")
    for run in (left, right):
        run.font.name = font_name
        run.font.size = Pt(spec["size_pt"])
        run.font.bold = spec["bold"]
        set_rfonts(run, font_name, font_name)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("eastAsia", "ascii", "hAnsi", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font_name)
    r_pr.append(r_fonts)
    if spec["bold"]:
        bold = OxmlElement("w:b")
        r_pr.append(bold)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(round(spec["size_pt"] * 2)))
    r_pr.append(size)
    field_run.insert(0, r_pr)


def set_footer(document: Document, profile: dict, font_name: str) -> None:
    duplex = profile["page"]["print_mode"] == "duplex"
    document.settings.odd_and_even_pages_header_footer = duplex
    page_spec = profile["page_number"]
    for section in document.sections:
        section.different_first_page_header_footer = False
        section.footer.is_linked_to_previous = False
        odd = section.footer.paragraphs[0]
        odd.alignment = WD_ALIGN_PARAGRAPH.RIGHT if duplex else WD_ALIGN_PARAGRAPH.CENTER
        add_page_field(odd, page_spec, font_name)
        if duplex:
            section.even_page_footer.is_linked_to_previous = False
            even = section.even_page_footer.paragraphs[0]
            even.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_page_field(even, page_spec, font_name)


def resolve_fonts(profile: dict) -> tuple[dict[str, str], str, list[str]]:
    installed = installed_font_names()
    warnings: list[str] = []
    fonts = {
        role: resolve_font(spec["font_cn"], spec["font_fallback"], installed, warnings)
        for role, spec in profile["styles"].items()
    }
    page = profile["page_number"]
    page_font = resolve_font(page["font_cn"], page["font_fallback"], installed, warnings)
    return fonts, page_font, warnings


def load_overrides(path: str | None) -> dict | None:
    return read_json(Path(path).expanduser().resolve()) if path else None


def load_classifications(path: str | None, document: Document) -> list[dict]:
    if not path:
        return classify_paragraphs(document)
    data = read_json(Path(path).expanduser().resolve())
    items = data.get("classifications")
    if not isinstance(items, list):
        raise ValueError("Classification file must contain a classifications array")
    seen: set[int] = set()
    for item in items:
        index = item.get("index")
        role = item.get("role")
        if not isinstance(index, int) or index < 0 or index >= len(document.paragraphs):
            raise ValueError(f"Invalid paragraph index: {index}")
        if role not in {*ROLE_NAMES, "skip"}:
            raise ValueError(f"Invalid role at paragraph {index}: {role}")
        if index in seen:
            raise ValueError(f"Duplicate paragraph index: {index}")
        seen.add(index)
    return items


def apply_document_format(document: Document, classifications: list[dict], profile: dict) -> list[str]:
    fonts, page_font, warnings = resolve_fonts(profile)
    ensure_styles(document, profile, fonts)
    for item in classifications:
        apply_role(document.paragraphs[item["index"]], item["role"], profile, fonts)
    apply_page_setup(document, profile, warnings)
    set_footer(document, profile, page_font)
    return warnings


def read_source_text(args: argparse.Namespace) -> str:
    if args.text is not None:
        return args.text
    return Path(args.input).expanduser().resolve().read_text(encoding="utf-8")


def new_output_path(output: str | None) -> Path:
    if output:
        result = Path(output).expanduser().resolve()
        if result.suffix.lower() != ".docx":
            raise ValueError("Output must use the .docx extension")
        return unique_output_path(result)
    return unique_output_path(Path.cwd() / "通用公文.docx")


def add_text(document: Document, text: str) -> list[dict]:
    explicit: dict[int, str] = {}
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        role = None
        if line.startswith("### "):
            role, line = "heading2", line[4:].strip()
        elif line.startswith("## "):
            role, line = "heading1", line[3:].strip()
        elif line.startswith("# "):
            role, line = "main_title", line[2:].strip()
        elif line.startswith("[说明]"):
            role, line = "description", line[4:].strip()
        elif line.startswith("[版记]"):
            role, line = "skip", line[4:].strip()
        paragraph = document.add_paragraph(line)
        if role:
            explicit[len(document.paragraphs) - 1] = role
    classifications = classify_paragraphs(document)
    for item in classifications:
        if item["index"] in explicit:
            item["role"] = explicit[item["index"]]
    return classifications


def finalize(document: Document, output: Path, profile: dict, classifications: list[dict]) -> dict:
    output.parent.mkdir(parents=True, exist_ok=True)
    warnings = apply_document_format(document, classifications, profile)
    document.save(output)
    scrub_docx(output)
    try:
        validation = validate_document(output, profile, classifications)
    except Exception:
        output.unlink(missing_ok=True)
        raise
    if not validation["valid"]:
        output.unlink(missing_ok=True)
        raise RuntimeError("Generated DOCX failed structural validation: " + "; ".join(validation["errors"]))
    return {"output": str(output), "warnings": sorted(set(warnings + validation["warnings"])), "validation": "structural_pass"}


def command_create(args: argparse.Namespace) -> None:
    profile = active_profile(load_overrides(args.overrides))
    if args.print_mode:
        profile["page"]["print_mode"] = args.print_mode
    document = Document()
    classifications = add_text(document, read_source_text(args))
    result = finalize(document, new_output_path(args.output), profile, classifications)
    print(json.dumps(result, ensure_ascii=False, indent=2))


def command_format_existing(args: argparse.Namespace) -> None:
    if not args.confirmed:
        raise SystemExit("Formatting requires --confirmed after the user reviews the inspection summary")
    source = Path(args.input).expanduser().resolve()
    inspection = inspect_document(source)
    if inspection.get("blockers"):
        raise SystemExit("; ".join(inspection["blockers"]))
    profile = active_profile(load_overrides(args.overrides))
    if args.print_mode:
        profile["page"]["print_mode"] = args.print_mode
    document = Document(source)
    classifications = load_classifications(args.classification, document)
    output = choose_output_path(source, args.output)
    result = finalize(document, output, profile, classifications)
    result["warnings"] = sorted(set(result["warnings"] + inspection.get("warnings", [])))
    print(json.dumps(result, ensure_ascii=False, indent=2))


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Create or reformat a general official-document DOCX")
    sub = parser.add_subparsers(dest="command", required=True)
    create = sub.add_parser("create")
    source = create.add_mutually_exclusive_group(required=True)
    source.add_argument("--input", help="UTF-8 text or Markdown file")
    source.add_argument("--text", help="Inline plain text or Markdown")
    create.add_argument("--output")
    create.add_argument("--overrides", help="Temporary JSON override file")
    create.add_argument("--print-mode", choices=["single", "duplex"])
    create.set_defaults(func=command_create)

    existing = sub.add_parser("format-existing")
    existing.add_argument("--input", required=True)
    existing.add_argument("--output")
    existing.add_argument("--classification", help="Reviewed classification JSON")
    existing.add_argument("--overrides", help="Temporary JSON override file")
    existing.add_argument("--print-mode", choices=["single", "duplex"])
    existing.add_argument("--confirmed", action="store_true")
    existing.set_defaults(func=command_format_existing)
    return parser


if __name__ == "__main__":
    arguments = build_parser().parse_args()
    try:
        arguments.func(arguments)
    except (KeyError, TypeError, ValueError) as exc:
        raise SystemExit(str(exc)) from exc
