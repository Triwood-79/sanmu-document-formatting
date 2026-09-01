from __future__ import annotations

import argparse
import json
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

from common import (
    active_profile,
    configure_utf8_stdio,
    read_json,
    signature_style_spec,
)
from inspect_docx import analyze_tables, classify_paragraphs, is_spacing_paragraph, signature_left_indents, split_signature_text, title_block_layout


STYLE_ROLES = {
    "ODF Main Title": "main_title",
    "ODF Heading 1": "heading1",
    "ODF Heading 2": "heading2",
    "ODF Body": "body",
    "ODF Reference Note": "reference_note",
    "ODF Description": "description",
    "ODF Signature": "signature",
    "ODF Colophon": "skip",
}
ROLE_STYLES = {role: style for style, role in STYLE_ROLES.items() if role not in {"colophon", "skip"}}
ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
FONT_ALIASES = {
    "黑体": {"SimHei"},
    "楷体_GB2312": {"楷体GB2312", "KaiTi_GB2312", "KaiTi"},
    "仿宋_GB2312": {"仿宋GB2312", "FangSong_GB2312", "FangSong"},
    "SimSun": {"宋体", "simsun"},
}


def near(actual: float | None, expected: float, tolerance: float = 0.05) -> bool:
    return actual is not None and abs(actual - expected) <= tolerance


def run_color_is_black(run) -> bool:
    r_pr = run._element.rPr
    color = r_pr.find(qn("w:color")) if r_pr is not None else None
    return color is not None and color.get(qn("w:val"), "").upper() == "000000"


def validate_run_font(run, label: str, spec: dict, errors: list[str]) -> bool:
    r_pr = run._element.rPr
    east_asia = r_pr.rFonts.get(qn("w:eastAsia")) if r_pr is not None and r_pr.rFonts is not None else None
    ascii_font = r_pr.rFonts.get(qn("w:ascii")) if r_pr is not None and r_pr.rFonts is not None else None
    hansi_font = r_pr.rFonts.get(qn("w:hAnsi")) if r_pr is not None and r_pr.rFonts is not None else None
    allowed_fonts = {spec["font_cn"], spec["font_fallback"]} | FONT_ALIASES.get(spec["font_fallback"], set())
    if east_asia not in allowed_fonts:
        errors.append(f"{label} has an unexpected East Asian font")
        return False
    if ascii_font != spec["font_latin"] or hansi_font != spec["font_latin"]:
        errors.append(f"{label} has an unexpected Latin font")
        return False
    if not run_color_is_black(run):
        errors.append(f"{label} is not black")
        return False
    return True


def validate_font_only(
    paragraph,
    label: str,
    spec: dict,
    errors: list[str],
    expected_bold: bool | None = None,
) -> None:
    validate_widow_control(paragraph, label, errors)
    for run in (run for run in paragraph.runs if run.text):
        if expected_bold is not None and run.font.bold is not expected_bold:
            errors.append(f"{label} has an incorrect bold setting")
            break
        if not validate_run_font(run, label, spec, errors):
            break


def validate_widow_control(paragraph, label: str, errors: list[str]) -> None:
    if paragraph.paragraph_format.widow_control is not False:
        errors.append(f"{label} must explicitly disable widow/orphan control")


def validate_black_only(paragraph, label: str, errors: list[str]) -> None:
    if any(run.text and not run_color_is_black(run) for run in paragraph.runs):
        errors.append(f"{label} is not black")


def validate_paragraph(paragraph, index: int, role: str, spec: dict, global_bold: bool, errors: list[str]) -> None:
    validate_widow_control(paragraph, f"Paragraph {index} ({role})", errors)
    p_pr = paragraph._p.pPr
    spacing = p_pr.spacing if p_pr is not None else None
    line = int(spacing.get(qn("w:line"))) / 20 if spacing is not None and spacing.get(qn("w:line")) else None
    line_rule = spacing.get(qn("w:lineRule")) if spacing is not None else None
    if line_rule != "exact" or not near(line, spec["line_spacing_pt"], 0.01):
        errors.append(f"Paragraph {index} ({role}) has incorrect exact line spacing")
    space_before = paragraph.paragraph_format.space_before
    before_pt = space_before.pt if space_before is not None else 0
    if not near(before_pt, spec["space_before_pt"], 0.01):
        errors.append(f"Paragraph {index} ({role}) has incorrect space before")
    if paragraph.alignment != ALIGNMENTS[spec["alignment"]]:
        errors.append(f"Paragraph {index} ({role}) has incorrect alignment")
    ind = p_pr.ind if p_pr is not None else None
    chars = int(ind.get(qn("w:firstLineChars"))) / 100 if ind is not None and ind.get(qn("w:firstLineChars")) else 0
    if not near(chars, spec["first_line_chars"], 0.01):
        errors.append(f"Paragraph {index} ({role}) has incorrect first-line character indent")
    text_runs = [run for run in paragraph.runs if run.text]
    for run in text_runs:
        size = run.font.size.pt if run.font.size else None
        if not near(size, spec["size_pt"], 0.01):
            errors.append(f"Paragraph {index} ({role}) has incorrect font size")
            break
        if run.font.bold is not global_bold:
            errors.append(f"Paragraph {index} ({role}) has incorrect bold setting")
            break
        if not validate_run_font(run, f"Paragraph {index} ({role})", spec, errors):
            break


def validate_signature_paragraph(paragraph, index: int, profile: dict, errors: list[str], expected_left: int) -> None:
    spec = signature_style_spec(profile)
    validate_paragraph(paragraph, index, "signature", spec, profile["global"]["bold"], errors)
    p_pr = paragraph._p.pPr
    ind = p_pr.ind if p_pr is not None else None
    left = int(ind.get(qn("w:left"))) if ind is not None and ind.get(qn("w:left")) else None
    left_chars = int(ind.get(qn("w:leftChars"))) if ind is not None and ind.get(qn("w:leftChars")) else None
    first_line = int(ind.get(qn("w:firstLine"))) if ind is not None and ind.get(qn("w:firstLine")) else 0
    if left != expected_left or left_chars is not None:
        errors.append(f"Paragraph {index} (signature) has incorrect right-side block positioning")
    if first_line != 0:
        errors.append(f"Paragraph {index} (signature) has an incorrect first-line indent")
    if split_signature_text(paragraph.text) is not None:
        errors.append(f"Paragraph {index} (signature) must separate the office and date into two paragraphs")
    if ind is not None and any(ind.get(qn(f"w:{attr}")) not in {None, "0"} for attr in ("right", "rightChars", "hanging", "hangingChars")):
        errors.append(f"Paragraph {index} (signature) has unexpected right or hanging indentation")


def validate_tables(document: Document, table_specs: list[dict], profile: dict, errors: list[str]) -> None:
    for table_spec in table_specs:
        table_index = table_spec["index"]
        title_index = table_spec.get("title_paragraph_index")
        if title_index is not None:
            validate_font_only(
                document.paragraphs[title_index],
                f"Table {table_index} title paragraph {title_index}",
                profile["styles"]["main_title"],
                errors,
                profile["global"]["bold"],
            )
        table = document.tables[table_index]
        header_rows = table_spec["header_rows"]
        for row_index, row in enumerate(table.rows):
            if row_index == 0 and header_rows >= 1:
                role = "heading1"
            elif row_index == 1 and header_rows >= 2:
                role = "heading2"
            else:
                role = "body"
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    validate_font_only(
                        paragraph,
                        f"Table {table_index} row {row_index} cell {cell_index} paragraph {paragraph_index}",
                        profile["styles"][role],
                        errors,
                        profile["global"]["bold"],
                    )


def validate_document(
    path: Path,
    profile: dict | None = None,
    classifications: list[dict] | None = None,
    table_specs: list[dict] | None = None,
) -> dict:
    path = path.expanduser().resolve()
    errors: list[str] = []
    warnings: list[str] = []
    if path.suffix.lower() != ".docx" or not zipfile.is_zipfile(path):
        return {"valid": False, "errors": ["Not a valid .docx package"], "warnings": []}
    profile = profile or active_profile()
    try:
        document = Document(path)
    except Exception as exc:
        return {"valid": False, "errors": [f"DOCX open failed: {type(exc).__name__}"], "warnings": []}

    margins = profile["page"]["margins_cm"]
    for number, section in enumerate(document.sections, start=1):
        if section.orientation == WD_ORIENT.LANDSCAPE:
            warnings.append(f"Landscape section {number} was preserved")
            continue
        values = {
            "top": section.top_margin.cm,
            "bottom": section.bottom_margin.cm,
            "left": section.left_margin.cm,
            "right": section.right_margin.cm,
        }
        for key, expected in margins.items():
            if not near(values[key], expected):
                errors.append(f"Section {number} has incorrect {key} margin")
        if not near(section.page_width.mm, 210, 0.1) or not near(section.page_height.mm, 297, 0.1):
            errors.append(f"Section {number} is not A4 portrait size")

    table_specs = table_specs if table_specs is not None else analyze_tables(document)
    table_title_indexes = {
        item["title_paragraph_index"] for item in table_specs if item.get("title_paragraph_index") is not None
    }
    expected_roles = (
        {item["index"]: item["role"] for item in classifications}
        if classifications is not None
        else {item["index"]: item["role"] for item in classify_paragraphs(document)}
    )
    title_indices = sorted(index for index, role in expected_roles.items() if role == "main_title")
    for index in title_indices:
        end_index, blanks = title_block_layout(document, expected_roles, index)
        spacer_index = end_index + 1
        if spacer_index not in blanks:
            errors.append(f"Title/description block at paragraph {index} is not followed by a blank line")
        elif expected_roles.get(spacer_index) != "body":
            errors.append(f"Title/description block at paragraph {index} blank line is not body-formatted")
        if any(blank != spacer_index for blank in blanks):
            errors.append(f"Title/description block at paragraph {index} has misplaced or duplicate blank lines")

    signature_content = sorted(
        index
        for index, role in expected_roles.items()
        if role == "signature" and index < len(document.paragraphs) and document.paragraphs[index].text.strip()
    )
    if signature_content:
        first_signature = signature_content[0]
        spacer_index = first_signature - 1
        if spacer_index < 0 or not is_spacing_paragraph(document.paragraphs[spacer_index]):
            errors.append(f"Paragraph {first_signature} (signature) is not preceded by a blank line")
        elif expected_roles.get(spacer_index) != "signature":
            errors.append(f"Paragraph {first_signature} (signature) blank line is not signature-formatted")
        if spacer_index > 0 and is_spacing_paragraph(document.paragraphs[spacer_index - 1]):
            if document.paragraphs[spacer_index - 1]._p.getnext() is document.paragraphs[spacer_index]._p:
                errors.append(f"Paragraph {first_signature} (signature) has more than one preceding blank line")

    signature_indents = signature_left_indents(document, profile, expected_roles)
    for index, paragraph in enumerate(document.paragraphs):
        if index in table_title_indexes:
            continue
        role = expected_roles.get(index)
        if role in {"skip", "colophon"}:
            continue
        elif role == "signature":
            if classifications is not None and paragraph.style.name != ROLE_STYLES[role]:
                errors.append(f"Paragraph {index} ({role}) is missing its managed style")
            validate_signature_paragraph(paragraph, index, profile, errors, signature_indents[index])
        elif role in profile["styles"]:
            if classifications is not None and paragraph.style.name != ROLE_STYLES[role]:
                errors.append(f"Paragraph {index} ({role}) is missing its managed style")
            validate_paragraph(
                paragraph,
                index,
                role,
                profile["styles"][role],
                profile["global"]["bold"],
                errors,
            )
        elif paragraph.text.strip():
            validate_black_only(paragraph, f"Paragraph {index}", errors)
            warnings.append(f"Paragraph {index} was preserved without managed formatting")

    validate_tables(document, table_specs, profile, errors)

    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        footer_parts = [archive.read(name) for name in names if name.startswith("word/footer") and name.endswith(".xml")]
        footer_xml = b"".join(footer_parts)
        if b"PAGE" not in footer_xml:
            errors.append("PAGE field is missing from the footer")
        page_spec = profile["page_number"]
        footer_alignments: set[str] = set()
        for data in footer_parts:
            root = etree.fromstring(data)
            field = root.find(".//w:fldSimple", namespaces={"w": W_NS})
            simple_instruction = field.get(qn("w:instr"), "") if field is not None else ""
            complex_instructions = root.xpath(".//w:instrText/text()", namespaces={"w": W_NS})
            if "PAGE" not in simple_instruction and not any("PAGE" in item for item in complex_instructions):
                continue
            text = "".join(root.xpath(".//w:t/text()", namespaces={"w": W_NS}))
            if text != "— 1 —":
                errors.append("Page-number decoration is incorrect")
            paragraph = root.find(".//w:p", namespaces={"w": W_NS})
            widow = paragraph.find("./w:pPr/w:widowControl", namespaces={"w": W_NS}) if paragraph is not None else None
            if widow is None or widow.get(qn("w:val"), "1") not in {"0", "false", "off"}:
                errors.append("Page-number paragraph must explicitly disable widow/orphan control")
            jc = paragraph.find("./w:pPr/w:jc", namespaces={"w": W_NS}) if paragraph is not None else None
            if jc is not None:
                footer_alignments.add(jc.get(qn("w:val"), ""))
            if field is not None:
                r_pr = field.find(".//w:rPr", namespaces={"w": W_NS})
            else:
                result_runs = root.xpath(
                    ".//w:r[w:t='1']",
                    namespaces={"w": W_NS},
                )
                r_pr = result_runs[0].find("./w:rPr", namespaces={"w": W_NS}) if result_runs else None
            r_fonts = r_pr.find("./w:rFonts", namespaces={"w": W_NS}) if r_pr is not None else None
            allowed_page_fonts = {page_spec["font_cn"], page_spec["font_fallback"]} | FONT_ALIASES.get(page_spec["font_fallback"], set())
            for font_slot in ("eastAsia", "ascii", "hAnsi", "cs"):
                actual_font = r_fonts.get(qn(f"w:{font_slot}")) if r_fonts is not None else None
                if actual_font not in allowed_page_fonts:
                    errors.append(f"Page number has an unexpected {font_slot} font")
            hint = r_pr.find("./w:hint", namespaces={"w": W_NS}) if r_pr is not None else None
            if hint is None or hint.get(qn("w:val")) != "eastAsia":
                errors.append("Page number must use the East Asian font hint")
            size = r_pr.find("./w:sz", namespaces={"w": W_NS}) if r_pr is not None else None
            actual_half_points = int(size.get(qn("w:val"))) if size is not None and size.get(qn("w:val")) else None
            if actual_half_points != round(page_spec["size_pt"] * 2):
                errors.append("Page number has an incorrect font size")
            bold = r_pr.find("./w:b", namespaces={"w": W_NS}) if r_pr is not None else None
            actual_bold = bold is not None and bold.get(qn("w:val"), "1") not in {"0", "false", "off"}
            if actual_bold is not page_spec["bold"]:
                errors.append("Page number has an incorrect bold setting")
            field_color = r_pr.find("./w:color", namespaces={"w": W_NS}) if r_pr is not None else None
            if field_color is None or field_color.get(qn("w:val"), "").upper() != "000000":
                errors.append("Page number is not black")
            for run in paragraph.findall(".//w:r", namespaces={"w": W_NS}) if paragraph is not None else []:
                if not run.findall(".//w:t", namespaces={"w": W_NS}):
                    continue
                run_r_pr = run.find("./w:rPr", namespaces={"w": W_NS})
                run_fonts = run_r_pr.find("./w:rFonts", namespaces={"w": W_NS}) if run_r_pr is not None else None
                for font_slot in ("eastAsia", "ascii", "hAnsi", "cs"):
                    actual_font = run_fonts.get(qn(f"w:{font_slot}")) if run_fonts is not None else None
                    if actual_font not in allowed_page_fonts:
                        errors.append(f"Page-number run has an unexpected {font_slot} font")
                        break
                run_hint = run_r_pr.find("./w:hint", namespaces={"w": W_NS}) if run_r_pr is not None else None
                if run_hint is None or run_hint.get(qn("w:val")) != "eastAsia":
                    errors.append("Page-number run must use the East Asian font hint")
                run_color = run.find("./w:rPr/w:color", namespaces={"w": W_NS})
                if run_color is None or run_color.get(qn("w:val"), "").upper() != "000000":
                    errors.append("Page-number decoration is not black")
                    break
        expected_alignments = {"right", "left"} if profile["page"]["print_mode"] == "duplex" else {"center"}
        if not expected_alignments.issubset(footer_alignments):
            errors.append("Page-number footer alignment is incorrect")
        settings = archive.read("word/settings.xml") if "word/settings.xml" in names else b""
        if profile["page"]["print_mode"] == "duplex" and b"evenAndOddHeaders" not in settings:
            errors.append("Duplex mode is missing even/odd footer settings")
        for metadata_name in ("docProps/core.xml", "docProps/app.xml", "docProps/custom.xml"):
            if metadata_name == "docProps/custom.xml" and metadata_name in names:
                errors.append("Custom document properties were not removed")
        core = archive.read("docProps/core.xml") if "docProps/core.xml" in names else b""
        if b"creator" in core or b"lastModifiedBy" in core:
            errors.append("Identifying core properties were not removed")

    return {"valid": not errors, "errors": errors, "warnings": warnings, "mode": "structural"}


def main() -> None:
    parser = argparse.ArgumentParser(description="Structurally validate a formatted DOCX")
    parser.add_argument("input")
    parser.add_argument("--profile", help="Optional complete profile JSON")
    args = parser.parse_args()
    profile = read_json(Path(args.profile).expanduser().resolve()) if args.profile else None
    result = validate_document(Path(args.input), profile)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    raise SystemExit(0 if result["valid"] else 1)


if __name__ == "__main__":
    configure_utf8_stdio()
    main()
