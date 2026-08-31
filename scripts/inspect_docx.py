from __future__ import annotations

import argparse
import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from common import configure_utf8_stdio, signature_indent_twips


H1_RE = re.compile(r"^[一二三四五六七八九十百]+、")
H2_RE = re.compile(r"^（[一二三四五六七八九十百]+）")
DATE_PATTERN = r"(?:\d{4}年\d{1,2}月\d{1,2}日|某年某月某日)"
DATE_RE = re.compile(rf"^{DATE_PATTERN}$")
SIGNATURE_DATE_RE = re.compile(rf"^{DATE_PATTERN}(?:\s*印发)?[。．.]?$")
INLINE_SIGNATURE_RE = re.compile(rf"^(.+?)\s*({DATE_PATTERN}(?:\s*印发)?[。．.]?)$", re.DOTALL)
NOTE_RE = re.compile(r"^(?:（.*）|\(.*\))$", re.DOTALL)
SIGNATURE_OFFICE_RE = re.compile(r"(?:单位|办公室|委员会|人民政府|政府|局|厅|部|院|中心|公司|集团|学校|协会|研究院)$")
ALLOWED_EXTENSIONS = {".docx"}
TABLE_TITLE_SUFFIXES = ("统计表", "一览表", "明细表", "安排表", "情况表", "汇总表", "清单")
STYLE_ROLES = {
    "ODF Main Title": "main_title",
    "ODF Heading 1": "heading1",
    "ODF Heading 2": "heading2",
    "ODF Body": "body",
    "ODF Reference Note": "reference_note",
    "ODF Description": "description",
    "ODF Signature": "signature",
    "ODF Colophon": "skip",  # Legacy style: no longer a formatting role.
}


def split_signature_text(text: str) -> tuple[str, str] | None:
    """Recognize an office/date line without discarding the date's suffix."""
    match = INLINE_SIGNATURE_RE.fullmatch(text.strip())
    if not match:
        return None
    office, date = (part.strip() for part in match.groups())
    if len(office) > 40 or not SIGNATURE_OFFICE_RE.search(office) or re.search(r"[：:。；;，,\r\n]", office):
        return None
    return office, date


def normalize_legacy_roles(document: Document, classifications: list[dict]) -> None:
    """Accept old maps without retaining the retired colophon formatter."""
    inferred = {item["index"]: item["role"] for item in classify_paragraphs(document)}
    for item in classifications:
        if item["role"] == "colophon":
            item["role"] = "signature" if inferred.get(item["index"]) == "signature" else "skip"


def signature_left_indents(document: Document, profile: dict, roles: dict[int, str]) -> dict[int, int]:
    """Use the actual section's text width; all signature lines share an axis."""
    lines = [p.text for index, p in enumerate(document.paragraphs) if roles.get(index) == "signature" and p.text.strip()]
    indents: dict[int, int] = {}
    section_index = 0
    sections = list(document.sections)
    for index, paragraph in enumerate(document.paragraphs):
        if roles.get(index) == "signature":
            section = sections[min(section_index, len(sections) - 1)]
            width = section.page_width.twips - section.left_margin.twips - section.right_margin.twips
            if section.gutter is not None:
                width -= section.gutter.twips
            indents[index] = signature_indent_twips(lines, profile["styles"]["body"]["size_pt"], width)
        if paragraph._p.pPr is not None and paragraph._p.pPr.find(qn("w:sectPr")) is not None:
            section_index += 1
    return indents


def is_spacing_paragraph(paragraph: Paragraph) -> bool:
    """Only plain empty paragraphs may be moved or removed as spacing."""
    if paragraph.text.strip():
        return False
    p_pr = paragraph._p.pPr
    if p_pr is not None and any(p_pr.find(qn(tag)) is not None for tag in ("w:sectPr", "w:pageBreakBefore")):
        return False
    for child in paragraph._p:
        if child.tag == qn("w:pPr"):
            continue
        if child.tag != qn("w:r"):
            return False
        if any(node.tag not in {qn("w:rPr"), qn("w:t")} for node in child):
            return False
    return True


def title_block_layout(document: Document, roles: dict[int, str], title_index: int) -> tuple[int, list[int]]:
    """Find the title/description block and its plain spacing paragraphs.

    Stop at the first other content block, including tables and drawings.
    Legacy blanks between the title and descriptions are included for repair.
    """
    paragraphs = document.paragraphs
    end_index = title_index
    blanks: list[int] = []
    previous = paragraphs[title_index]._p
    for index in range(title_index + 1, len(paragraphs)):
        paragraph = paragraphs[index]
        if previous.getnext() is not paragraph._p:
            break
        if is_spacing_paragraph(paragraph) and roles.get(index) in {None, "skip", "body", "description"}:
            blanks.append(index)
        elif paragraph.text.strip() and roles.get(index) == "description":
            end_index = index
        else:
            break
        previous = paragraph._p
    return end_index, blanks


def package_risks(path: Path) -> tuple[list[str], list[str]]:
    blockers: list[str] = []
    warnings: list[str] = []
    if path.suffix.lower() not in ALLOWED_EXTENSIONS:
        blockers.append("Only .docx files are supported; convert legacy or macro-enabled files first")
        return blockers, warnings
    if not zipfile.is_zipfile(path):
        blockers.append("The file is corrupt, encrypted, or not a valid DOCX package")
        return blockers, warnings
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        if "word/comments.xml" in names or any(name.startswith("word/comments") for name in names):
            blockers.append("Comments are present; provide a clean copy without comments")
        document_xml = archive.read("word/document.xml") if "word/document.xml" in names else b""
        if b"<w:ins" in document_xml or b"<w:del" in document_xml or b"<w:moveFrom" in document_xml:
            blockers.append("Tracked changes are present; accept or reject them before formatting")
        if b"<w:documentProtection" in archive.read("word/settings.xml") if "word/settings.xml" in names else False:
            blockers.append("Document protection is enabled")
        if "word/header1.xml" in names or "word/footer1.xml" in names:
            warnings.append("Existing header or footer content may be replaced when standard page numbers are applied")
    return blockers, warnings


def classify_paragraphs(document: Document) -> list[dict]:
    entries: list[dict] = []
    nonempty = [index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip()]
    title_index = nonempty[0] if nonempty else None
    description_indices: set[int] = set()
    signature_indices: set[int] = set()
    if title_index is not None:
        after = [index for index in nonempty if index > title_index][:2]
        if after:
            if DATE_RE.fullmatch(document.paragraphs[after[0]].text.strip()):
                description_indices.add(after[0])
            elif len(after) > 1 and DATE_RE.fullmatch(document.paragraphs[after[1]].text.strip()):
                if len(document.paragraphs[after[0]].text.strip()) <= 40:
                    description_indices.add(after[0])
                description_indices.add(after[1])
            else:
                office_text = document.paragraphs[after[0]].text.strip()
                if (
                    len(office_text) <= 40
                    and SIGNATURE_OFFICE_RE.search(office_text)
                    and not H1_RE.match(office_text)
                    and not H2_RE.match(office_text)
                    and not NOTE_RE.fullmatch(office_text)
                ):
                    description_indices.add(after[0])
    if nonempty and nonempty[-1] != title_index and split_signature_text(document.paragraphs[nonempty[-1]].text):
        signature_indices.add(nonempty[-1])
    if len(nonempty) >= 2:
        date_index = nonempty[-1]
        office_index = nonempty[-2]
        office_text = document.paragraphs[office_index].text.strip()
        if (
            SIGNATURE_DATE_RE.fullmatch(document.paragraphs[date_index].text.strip())
            and office_index != title_index
            and len(office_text) <= 40
            and SIGNATURE_OFFICE_RE.search(office_text)
            and not H1_RE.match(office_text)
            and not H2_RE.match(office_text)
            and not NOTE_RE.fullmatch(office_text)
        ):
            signature_indices.update({office_index, date_index})
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        styled_role = STYLE_ROLES.get(paragraph.style.name)
        if index in signature_indices and styled_role in {None, "body", "signature", "skip"}:
            role = "signature"
        elif styled_role:
            role = styled_role
        elif not text:
            role = "skip"
        elif index == title_index:
            role = "main_title"
        elif index in description_indices:
            role = "description"
        elif H1_RE.match(text):
            role = "heading1"
        elif H2_RE.match(text):
            role = "heading2"
        elif NOTE_RE.fullmatch(text):
            role = "reference_note"
        else:
            role = "body"
        entries.append({"index": index, "role": role, "text_preview": text[:80]})
    return entries


def looks_like_table_title(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    if not text or len(text) > 80 or text.endswith(("。", "；", "：", ":")):
        return False
    explicit_number = re.match(r"^(?:附?表)\s*[一二三四五六七八九十百\d]+", text) is not None
    descriptive_suffix = text.endswith(TABLE_TITLE_SUFFIXES)
    centered_table_name = paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER and text.endswith("表")
    return explicit_number or descriptive_suffix or centered_table_name


def row_is_repeating_header(row) -> bool:
    tr_pr = row._tr.trPr
    return tr_pr is not None and tr_pr.find(qn("w:tblHeader")) is not None


def row_has_merge(row) -> bool:
    cells = row.cells
    if len({id(cell._tc) for cell in cells}) < len(cells):
        return True
    for cell in cells:
        tc_pr = cell._tc.tcPr
        if tc_pr is not None and tc_pr.find(qn("w:vMerge")) is not None:
            return True
    return False


def analyze_tables(document: Document) -> list[dict]:
    blocks = list(document.iter_inner_content())
    tables: list[dict] = []
    paragraph_index = -1
    table_index = -1
    for block_index, block in enumerate(blocks):
        if isinstance(block, Paragraph):
            paragraph_index += 1
            continue
        if not isinstance(block, Table):
            continue
        table_index += 1
        title_index = None
        title_preview = None
        previous = blocks[block_index - 1] if block_index else None
        if isinstance(previous, Paragraph) and looks_like_table_title(previous):
            title_index = paragraph_index
            title_preview = previous.text.strip()[:80]
        header_rows = 1 if block.rows else 0
        if len(block.rows) > 1 and (row_is_repeating_header(block.rows[1]) or row_has_merge(block.rows[0])):
            header_rows = 2
        tables.append({
            "index": table_index,
            "title_paragraph_index": title_index,
            "title_preview": title_preview,
            "header_rows": header_rows,
            "rows": len(block.rows),
            "columns": len(block.columns),
        })
    return tables


def inspect_document(path: Path) -> dict:
    blockers, warnings = package_risks(path)
    if blockers and (path.suffix.lower() != ".docx" or not zipfile.is_zipfile(path)):
        return {"input": path.name, "blockers": blockers, "warnings": warnings}
    try:
        document = Document(path)
    except Exception as exc:
        blockers.append(f"DOCX could not be opened: {type(exc).__name__}")
        return {"input": path.name, "blockers": blockers, "warnings": warnings}

    drawing_count = 0
    textbox_count = 0
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml")
        drawing_count = xml.count(b"<w:drawing") + xml.count(b"<w:pict")
        textbox_count = xml.count(b"<w:txbxContent")
    landscape_sections = sum(1 for section in document.sections if section.orientation == WD_ORIENT.LANDSCAPE)
    classifications = classify_paragraphs(document)
    if landscape_sections:
        warnings.append("Landscape sections will be preserved and will not receive portrait page settings")
    if document.tables:
        warnings.append("Table fonts, text color, and the global bold setting will be normalized; borders, shading, alignment, sizes, spacing, row heights, and column widths will be preserved")
    if drawing_count or textbox_count:
        warnings.append("Complex elements will be preserved without internal reformatting")

    tables = analyze_tables(document)
    counts: dict[str, int] = {}
    for item in classifications:
        counts[item["role"]] = counts.get(item["role"], 0) + 1
    return {
        "input": path.name,
        "blockers": blockers,
        "warnings": warnings,
        "summary": {
            "top_level_paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
            "table_title_candidates": sum(1 for table in tables if table["title_paragraph_index"] is not None),
            "drawings": drawing_count,
            "textboxes": textbox_count,
            "landscape_sections": landscape_sections,
            "classification_counts": counts,
        },
        "classifications": classifications,
        "tables": tables,
        "confirmation_required": not blockers,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Inspect and classify a DOCX before formatting")
    parser.add_argument("input")
    parser.add_argument("--write-map", help="Write the classification JSON to this path")
    args = parser.parse_args()
    result = inspect_document(Path(args.input).expanduser().resolve())
    if args.write_map and "classifications" in result:
        output = Path(args.write_map).expanduser().resolve()
        output.write_text(
            json.dumps({"classifications": result["classifications"], "tables": result.get("tables", [])}, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    raise SystemExit(2 if result.get("blockers") else 0)


if __name__ == "__main__":
    configure_utf8_stdio()
    main()
