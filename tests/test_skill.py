from __future__ import annotations

import base64
import io
import json
import os
import subprocess
import sys
import tempfile
import unittest
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS))

from common import active_profile, load_preset, signature_indent_twips  # noqa: E402
from inspect_docx import analyze_tables, classify_paragraphs, inspect_document  # noqa: E402
from docx_engine import add_text, ensure_required_spacing, finalize  # noqa: E402
from privacy_scrub import scrub_docx  # noqa: E402
from validate_docx import validate_document  # noqa: E402


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Wl2nWQAAAAASUVORK5CYII="
)


def run_east_asia_font(run) -> str | None:
    r_pr = run._element.rPr
    return r_pr.rFonts.get(qn("w:eastAsia")) if r_pr is not None and r_pr.rFonts is not None else None


def run_color(run) -> str | None:
    r_pr = run._element.rPr
    color = r_pr.find(qn("w:color")) if r_pr is not None else None
    return color.get(qn("w:val")) if color is not None else None


class SkillTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory(prefix="odf-test-")
        self.root = Path(self.temp.name)
        self.env = os.environ.copy()
        self.env["CODEX_HOME"] = str(self.root / "codex-state")
        self.env["PYTHONIOENCODING"] = "utf-8"

    def tearDown(self) -> None:
        self.temp.cleanup()

    def run_script(self, name: str, *args: str, check: bool = True) -> subprocess.CompletedProcess:
        return subprocess.run(
            [sys.executable, str(SCRIPTS / name), *map(str, args)],
            env=self.env,
            capture_output=True,
            text=True,
            encoding="utf-8",
            check=check,
        )

    def test_preset_exact_values(self) -> None:
        profile = load_preset()
        self.assertEqual(profile["page"]["margins_cm"], {"top": 3.7, "bottom": 3.5, "left": 2.8, "right": 2.6})
        self.assertEqual(profile["page"]["print_mode"], "duplex")
        self.assertEqual(profile["styles"]["main_title"]["size_pt"], 22)
        self.assertEqual(profile["styles"]["reference_note"]["size_pt"], 15)
        self.assertEqual(profile["page_number"]["size_pt"], 14)
        self.assertTrue(profile["global"]["bold"])
        self.assertNotIn("colophon", profile["styles"])
        self.assertNotIn("signature", profile["styles"])

    def test_cli_forces_utf8_output_even_when_python_defaults_to_gbk(self) -> None:
        env = self.env.copy()
        env["PYTHONIOENCODING"] = "gbk"
        result = subprocess.run(
            [sys.executable, str(SCRIPTS / "profile_manager.py"), "view"],
            env=env,
            capture_output=True,
            check=True,
        )
        output = result.stdout.decode("utf-8")
        self.assertIn("宋体", output)
        self.assertEqual(json.loads(output)["profile"]["page_number"]["font_cn"], "宋体")

    def test_profile_draft_confirmation_and_restore(self) -> None:
        self.run_script("profile_manager.py", "begin", "--fresh")
        self.run_script("profile_manager.py", "set", "page.margins_cm.top", "4.0")
        diff = json.loads(self.run_script("profile_manager.py", "diff").stdout)
        self.assertEqual(diff["changes"][0]["current"], 3.7)
        self.assertEqual(diff["changes"][0]["proposed"], 4.0)
        before = json.loads(self.run_script("profile_manager.py", "view").stdout)
        self.assertEqual(before["profile"]["page"]["margins_cm"]["top"], 3.7)
        self.run_script("profile_manager.py", "confirm")
        after = json.loads(self.run_script("profile_manager.py", "view").stdout)
        self.assertEqual(after["profile"]["page"]["margins_cm"]["top"], 4.0)
        self.run_script("profile_manager.py", "restore")
        restored = json.loads(self.run_script("profile_manager.py", "view").stdout)
        self.assertEqual(restored["source"], "built_in_preset")

    def test_invalid_and_fixed_profile_fields_are_rejected(self) -> None:
        self.run_script("profile_manager.py", "begin", "--fresh")
        invalid_alignment = self.run_script(
            "profile_manager.py", "set", "styles.body.alignment", '"diagonal"', check=False
        )
        self.assertNotEqual(invalid_alignment.returncode, 0)
        self.assertIn("Unsupported value", invalid_alignment.stderr)
        fixed_field = self.run_script(
            "profile_manager.py", "set", "page_number.decoration", '"none"', check=False
        )
        self.assertNotEqual(fixed_field.returncode, 0)
        self.assertIn("fixed in V1", fixed_field.stderr)

    def test_page_number_bold_can_be_disabled_and_validated(self) -> None:
        self.run_script("profile_manager.py", "begin", "--fresh")
        self.run_script("profile_manager.py", "set", "page_number.bold", "false")
        self.run_script("profile_manager.py", "confirm")
        output = self.root / "page-number-not-bold.docx"
        self.run_script("docx_engine.py", "create", "--text", "# 示例标题\n正文。", "--output", output)
        profile = load_preset()
        profile["page_number"]["bold"] = False
        report = validate_document(output, profile)
        self.assertTrue(report["valid"], report["errors"])

    def test_create_and_validate_duplex_document(self) -> None:
        output = self.root / "generated.docx"
        text = "# 示例标题\n[说明]某单位\n[说明]某年某月某日\n## 一、一级标题\n正文 Test 123。\n（补充参考资料。）"
        result = json.loads(self.run_script("docx_engine.py", "create", "--text", text, "--output", output).stdout)
        self.assertEqual(result["validation"], "structural_pass")
        report = validate_document(output, load_preset())
        self.assertTrue(report["valid"], report["errors"])
        with zipfile.ZipFile(output) as archive:
            settings = archive.read("word/settings.xml")
            footers = [archive.read(name) for name in archive.namelist() if name.startswith("word/footer")]
        self.assertIn(b"evenAndOddHeaders", settings)
        self.assertTrue(any(b"PAGE" in footer for footer in footers))
        self.assertTrue(any(b'w:jc w:val="right"' in footer for footer in footers))
        self.assertTrue(any(b'w:jc w:val="left"' in footer for footer in footers))

    def test_create_inserts_title_blank_line_and_formats_signature_block(self) -> None:
        output = self.root / "signature-created.docx"
        text = "# 示例公文标题\n各相关部门：\n这是正文。\n某测试单位办公室\n2026年8月31日"
        result = json.loads(self.run_script("docx_engine.py", "create", "--text", text, "--output", output).stdout)
        self.assertEqual(result["validation"], "structural_pass")

        document = Document(output)
        self.assertEqual(document.paragraphs[1].text, "")
        self.assertEqual(document.paragraphs[1].style.name, "ODF Body")
        self.assertEqual(document.paragraphs[-3].text, "")
        self.assertTrue(all(paragraph.style.name == "ODF Signature" for paragraph in document.paragraphs[-3:]))
        self.assertEqual(document.paragraphs[-2].text, "某测试单位办公室")
        self.assertEqual(document.paragraphs[-1].text, "2026年8月31日")
        for paragraph in document.paragraphs[-3:]:
            ind = paragraph._p.pPr.ind
            self.assertGreater(int(ind.get(qn("w:left"))), 0)
            self.assertIsNone(ind.get(qn("w:leftChars")))
            self.assertEqual(ind.get(qn("w:firstLine")), "0")
            self.assertEqual(ind.get(qn("w:firstLineChars")), "0")
            self.assertEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        report = validate_document(output, load_preset())
        self.assertTrue(report["valid"], report["errors"])

    def test_existing_document_adds_required_blank_lines_and_signature_geometry(self) -> None:
        source = self.root / "signature-source.docx"
        document = Document()
        document.add_paragraph("示例公文标题")
        document.add_paragraph("这是正文。")
        document.add_paragraph("某测试单位办公室")
        document.add_paragraph("2026年8月31日")
        document.save(source)

        inspection = inspect_document(source)
        self.assertEqual([item["role"] for item in inspection["classifications"][-2:]], ["signature", "signature"])
        mapping = self.root / "signature-map.json"
        mapping.write_text(
            json.dumps({"classifications": inspection["classifications"], "tables": []}, ensure_ascii=False),
            encoding="utf-8",
        )
        output = self.root / "signature-formatted.docx"
        self.run_script(
            "docx_engine.py",
            "format-existing",
            "--input",
            source,
            "--output",
            output,
            "--classification",
            mapping,
            "--confirmed",
        )

        formatted = Document(output)
        self.assertEqual([paragraph.text for paragraph in formatted.paragraphs], [
            "示例公文标题",
            "",
            "这是正文。",
            "",
            "某测试单位办公室",
            "2026年8月31日",
        ])
        self.assertTrue(all(paragraph.style.name == "ODF Signature" for paragraph in formatted.paragraphs[-3:]))
        self.assertEqual([paragraph.text for paragraph in Document(source).paragraphs], [
            "示例公文标题",
            "这是正文。",
            "某测试单位办公室",
            "2026年8月31日",
        ])

    def test_validator_rejects_missing_title_blank_line_and_bad_signature_position(self) -> None:
        output = self.root / "signature-invalid.docx"
        text = "# 示例公文标题\n这是正文。\n[落款]某测试单位办公室\n[落款]2026年8月31日"
        self.run_script("docx_engine.py", "create", "--text", text, "--output", output)
        document = Document(output)
        document.paragraphs[1].text = "不应出现的文字"
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        document.save(output)
        scrub_docx(output)

        report = validate_document(output, load_preset())
        self.assertFalse(report["valid"])
        self.assertTrue(any("not followed by a blank line" in error for error in report["errors"]))
        self.assertTrue(any("incorrect alignment" in error for error in report["errors"]))

    def test_existing_output_path_gets_numeric_suffix(self) -> None:
        existing = self.root / "named-output.docx"
        original = Document()
        original.add_paragraph("SENTINEL")
        original.save(existing)
        result = json.loads(
            self.run_script("docx_engine.py", "create", "--text", "# 新标题\n新正文。", "--output", existing).stdout
        )
        actual = Path(result["output"])
        self.assertEqual(actual.name, "named-output_2.docx")
        self.assertEqual(Document(existing).paragraphs[0].text, "SENTINEL")
        self.assertTrue(actual.exists())

    def test_create_explicitly_disables_widow_control_on_all_managed_roles(self) -> None:
        document = Document()
        classifications = add_text(document, "示例标题\n[说明]某署名\n一、示例标题\n（一）示例子标题\n正文。\n（示例备注）\n某单位办公室\n某年某月某日")
        output = self.root / "no-widow.docx"
        finalize(document, output, load_preset(), classifications, [])
        formatted = Document(output)
        for paragraph in formatted.paragraphs:
            self.assertIs(paragraph.paragraph_format.widow_control, False)
        for style in formatted.styles:
            if style.name.startswith("ODF "):
                self.assertIs(style.paragraph_format.widow_control, False)
        for section in formatted.sections:
            self.assertIs(section.footer.paragraphs[0].paragraph_format.widow_control, False)
            self.assertIs(section.even_page_footer.paragraphs[0].paragraph_format.widow_control, False)

    def test_existing_widow_control_is_disabled_without_changing_other_pagination(self) -> None:
        document = Document()
        document.styles["Normal"].paragraph_format.widow_control = True
        document.add_paragraph("示例标题")
        body = document.add_paragraph("正文。")
        body.paragraph_format.widow_control = True
        body.paragraph_format.keep_with_next = True
        body.paragraph_format.keep_together = True
        body.paragraph_format.page_break_before = True
        title = document.add_paragraph("示例统计表")
        table = document.add_table(rows=2, cols=1)
        for row in table.rows:
            paragraph = row.cells[0].paragraphs[0]
            paragraph.add_run("示例")
            paragraph.paragraph_format.widow_control = True
            paragraph.paragraph_format.keep_together = True
        preserved = document.add_paragraph("保留段落。")
        preserved.paragraph_format.widow_control = True
        preserved_xml = preserved._p.xml
        classifications = classify_paragraphs(document)
        classifications[-1]["role"] = "skip"
        output = self.root / "existing-no-widow.docx"
        finalize(document, output, load_preset(), classifications, analyze_tables(document))
        self.assertIs(body.paragraph_format.widow_control, False)
        self.assertIs(title.paragraph_format.widow_control, False)
        self.assertIs(body.paragraph_format.keep_with_next, True)
        self.assertIs(body.paragraph_format.keep_together, True)
        self.assertIs(body.paragraph_format.page_break_before, True)
        for row in table.rows:
            paragraph = row.cells[0].paragraphs[0]
            self.assertIs(paragraph.paragraph_format.widow_control, False)
            self.assertIs(paragraph.paragraph_format.keep_together, True)
        self.assertEqual(preserved._p.xml, preserved_xml)

    def test_validator_rejects_enabled_or_unspecified_widow_control(self) -> None:
        document = Document()
        document.add_paragraph("示例标题")
        document.add_paragraph("正文。")
        document.add_table(rows=1, cols=1).cell(0, 0).text = "表格示例"
        valid = self.root / "widow-valid.docx"
        finalize(document, valid, load_preset(), classify_paragraphs(document), analyze_tables(document))
        for target in ("body", "table", "footer"):
            for value in (True, None):
                with self.subTest(target=target, value=value):
                    changed = Document(valid)
                    if target == "body":
                        paragraph = next(p for p in changed.paragraphs if p.text == "正文。")
                    elif target == "table":
                        paragraph = changed.tables[0].cell(0, 0).paragraphs[0]
                    else:
                        paragraph = changed.sections[0].footer.paragraphs[0]
                    paragraph.paragraph_format.widow_control = value
                    invalid = self.root / "widow-invalid.docx"
                    changed.save(invalid)
                    scrub_docx(invalid)
                    report = validate_document(invalid, load_preset())
                    self.assertFalse(report["valid"])
                    self.assertTrue(any("widow/orphan control" in error for error in report["errors"]))

    def test_header_spacing_optional_description_and_addressee_matrix(self) -> None:
        for descriptions in ([], ["某示例单位"], ["某年某月某日"], ["某示例单位", "某年某月某日"]):
            for addressee in ([], ["各相关处室："]):
                lines = ["示例标题", *descriptions, *addressee, "为做好示例工作，现通知如下。"]
                expected = ["示例标题", *descriptions, "", *addressee, lines[-1]]
                for mode in ("create", "existing"):
                    with self.subTest(descriptions=descriptions, addressee=addressee, mode=mode):
                        document = Document()
                        if mode == "create":
                            classifications = add_text(document, "\n".join(lines))
                        else:
                            for line in lines:
                                document.add_paragraph(line)
                            classifications = classify_paragraphs(document)
                        output = self.root / "header-matrix.docx"
                        finalize(document, output, load_preset(), classifications, [])
                        formatted = Document(output)
                        self.assertEqual([p.text for p in formatted.paragraphs], expected)
                        spacer = formatted.paragraphs[1 + len(descriptions)]
                        self.assertEqual(spacer.style.name, "ODF Body")
                        self.assertEqual(spacer._p.pPr.spacing.get(qn("w:line")), "600")
                        self.assertEqual(spacer._p.pPr.spacing.get(qn("w:lineRule")), "exact")
                        report = validate_document(output, load_preset())
                        self.assertTrue(report["valid"], report["errors"])

    def test_header_spacing_repair_preserves_maps_tables_and_is_idempotent(self) -> None:
        for blanks in ((1, 0, 0), (1, 1, 1), (0, 0, 2)):
            with self.subTest(blanks=blanks):
                document = Document()
                for line, count in zip(["示例标题", "某示例单位", "某年某月某日"], blanks):
                    document.add_paragraph(line)
                    for _ in range(count):
                        document.add_paragraph("")
                for line in ["各相关处室：", "这是正文。", "示例数据统计表"]:
                    document.add_paragraph(line)
                document.add_table(rows=2, cols=2)
                document.add_paragraph("某示例单位办公室")
                document.add_paragraph("某年某月某日")
                source = self.root / "legacy-source.docx"
                document.save(source)
                source_bytes = source.read_bytes()
                classifications = classify_paragraphs(document)
                tables = analyze_tables(document)
                self.assertIsNotNone(tables[0]["title_paragraph_index"])
                output = self.root / "legacy-repaired.docx"
                finalize(document, output, load_preset(), classifications, tables)
                expected = [
                    "示例标题", "某示例单位", "某年某月某日", "", "各相关处室：", "这是正文。",
                    "示例数据统计表", "", "某示例单位办公室", "某年某月某日",
                ]
                self.assertEqual([p.text for p in document.paragraphs], expected)
                self.assertEqual(tables[0]["title_paragraph_index"], 6)
                for item in classifications:
                    self.assertEqual(document.paragraphs[item["index"]].text[:80], item["text_preview"])
                self.assertEqual(source.read_bytes(), source_bytes)
                for _ in range(2):
                    document = Document(output)
                    classifications = classify_paragraphs(document)
                    tables = analyze_tables(document)
                    finalize(document, output, load_preset(), classifications, tables)
                    self.assertEqual([p.text for p in Document(output).paragraphs], expected)
                report = validate_document(output, load_preset())
                self.assertTrue(report["valid"], report["errors"])

    def test_validator_rejects_legacy_misplaced_and_duplicate_header_blanks(self) -> None:
        document = Document()
        classifications = add_text(document, "示例标题\n某示例单位\n某年某月某日\n各相关处室：\n正文。")
        valid = self.root / "header-valid.docx"
        finalize(document, valid, load_preset(), classifications, [])
        for variant in ("legacy", "both", "duplicate"):
            with self.subTest(variant=variant):
                document = Document(valid)
                spacer = document.paragraphs[3]._p
                if variant == "legacy":
                    document.paragraphs[0]._p.addnext(spacer)
                else:
                    anchor = document.paragraphs[0]._p if variant == "both" else spacer
                    anchor.addnext(deepcopy(spacer))
                output = self.root / "header-invalid.docx"
                document.save(output)
                scrub_docx(output)
                report = validate_document(output, load_preset())
                self.assertFalse(report["valid"])
                self.assertTrue(any("misplaced or duplicate" in error for error in report["errors"]))
                if variant == "legacy":
                    self.assertTrue(any("not followed by a blank line" in error for error in report["errors"]))

    def test_header_spacing_preserves_nontext_content_and_stops_at_tables(self) -> None:
        for kind in ("drawing", "break", "bookmark", "section", "table"):
            with self.subTest(kind=kind):
                document = Document()
                title = document.add_paragraph("示例标题")
                if kind == "table":
                    block = document.add_table(rows=1, cols=1)._tbl
                else:
                    paragraph = document.add_paragraph()
                    block = paragraph._p
                    if kind == "drawing":
                        paragraph.add_run().add_picture(io.BytesIO(PNG_1X1))
                    elif kind == "break":
                        paragraph.add_run().add_break()
                    elif kind == "bookmark":
                        mark = OxmlElement("w:bookmarkStart")
                        mark.set(qn("w:id"), "0")
                        mark.set(qn("w:name"), "example")
                        block.append(mark)
                    else:
                        block.get_or_add_pPr().append(OxmlElement("w:sectPr"))
                document.add_paragraph("某示例单位")
                document.add_paragraph("正文。")
                before = block.xml
                classifications = classify_paragraphs(document)
                ensure_required_spacing(document, classifications, analyze_tables(document))
                self.assertEqual(block.xml, before)
                self.assertIs(title._p.getnext().getnext(), block)
                self.assertEqual(document.paragraphs[1].style.name, "Normal")
                self.assertEqual(next(item for item in classifications if item["index"] == 1)["role"], "body")

    def test_explicit_description_without_date_is_kept_above_header_blank(self) -> None:
        document = Document()
        classifications = add_text(document, "# 示例标题\n[说明]某署名\n正文。")
        output = self.root / "explicit-description.docx"
        finalize(document, output, load_preset(), classifications, [])
        self.assertEqual([p.text for p in Document(output).paragraphs], ["示例标题", "某署名", "", "正文。"])

    def test_legacy_marker_migrates_terminal_office_date_to_signature(self) -> None:
        output = self.root / "colophon-unsupported.docx"
        result = json.loads(
            self.run_script(
                "docx_engine.py",
                "create",
                "--text",
                "# 示例标题\n正文。\n[版记]某单位办公室 某年某月某日印发",
                "--output",
                output,
            ).stdout
        )
        self.assertFalse(any("Colophon" in warning for warning in result["warnings"]))
        document = Document(output)
        self.assertEqual([p.text for p in document.paragraphs[-3:]], ["", "某单位办公室", "某年某月某日印发"])
        self.assertEqual(document.paragraphs[-1].style.name, "ODF Signature")
        self.assertIn(
            run_east_asia_font(document.paragraphs[-1].runs[0]),
            {"方正仿宋_GBK", "仿宋_GB2312", "仿宋GB2312", "FangSong_GB2312", "FangSong"},
        )
        self.assertEqual(run_color(document.paragraphs[-1].runs[0]), "000000")
        report = validate_document(output, load_preset())
        self.assertTrue(report["valid"], report["errors"])
        self.assertFalse(any("Colophon" in warning for warning in report["warnings"]))

    def test_existing_combined_office_date_is_split_and_formatted_as_signature(self) -> None:
        source = self.root / "existing-colophon.docx"
        document = Document()
        document.add_paragraph("示例标题")
        document.add_paragraph("这是正文。")
        colophon = document.add_paragraph()
        colophon.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = colophon.add_run("某测试单位办公室 2026年8月28日印发")
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(31, 78, 121)
        document.save(source)

        inspection = inspect_document(source)
        self.assertEqual(inspection["classifications"][-1]["role"], "signature")
        mapping = self.root / "existing-colophon-map.json"
        mapping.write_text(
            json.dumps({"classifications": inspection["classifications"], "tables": []}, ensure_ascii=False),
            encoding="utf-8",
        )
        output = self.root / "existing-colophon-formatted.docx"
        self.run_script(
            "docx_engine.py",
            "format-existing",
            "--input",
            source,
            "--output",
            output,
            "--classification",
            mapping,
            "--confirmed",
        )

        formatted = Document(output)
        paragraph = formatted.paragraphs[-1]
        formatted_run = paragraph.runs[0]
        self.assertEqual([p.text for p in formatted.paragraphs[-3:]], ["", "某测试单位办公室", "2026年8月28日印发"])
        self.assertEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(formatted_run.font.size.pt, 16)
        self.assertIn(
            run_east_asia_font(formatted_run),
            {"方正仿宋_GBK", "仿宋_GB2312", "仿宋GB2312", "FangSong_GB2312", "FangSong"},
        )
        self.assertEqual(run_color(formatted_run), "000000")

    def test_signature_indent_changes_with_text_size_and_page_width(self) -> None:
        lefts = []
        for office, size, margin in [("某单位", 16, 2.8), ("某示例联合协调工作委员会办公室", 16, 2.8),
                                     ("某示例联合协调工作委员会办公室", 20, 2.8), ("某单位", 16, 4.0)]:
            document = Document()
            profile = load_preset()
            profile["styles"]["body"]["size_pt"] = size
            profile["page"]["margins_cm"]["left"] = margin
            classifications = add_text(document, f"示例标题\n正文。\n[落款]{office}\n[落款]2026年8月28日印发")
            output = self.root / "dynamic-signature.docx"
            finalize(document, output, profile, classifications, [])
            formatted = Document(output)
            indents = [int(p._p.pPr.ind.get(qn("w:left"))) for p in formatted.paragraphs[-3:]]
            self.assertEqual(len(set(indents)), 1)
            lefts.append(indents[0])
            section = formatted.sections[0]
            available = section.page_width.twips - section.left_margin.twips - section.right_margin.twips
            # Reserve at least the full-width character budget for either line.
            required = max(len(office), len("2026年8月28日印发")) * size * 20
            self.assertGreaterEqual(available - indents[0], min(available, required))
            self.assertTrue(validate_document(output, profile)["valid"])
        self.assertLess(lefts[1], lefts[0])
        self.assertLess(lefts[2], lefts[1])
        self.assertLess(lefts[3], lefts[0])

    def test_signature_width_regression_for_mixed_digit_date(self) -> None:
        document = Document()
        date = "2026年8月28日印发"
        classifications = add_text(document, f"示例标题\n正文。\n某测试单位办公室\n{date}")
        output = self.root / "signature-width-regression.docx"
        finalize(document, output, load_preset(), classifications, [])
        formatted = Document(output)
        section = formatted.sections[-1]
        available = section.page_width.twips - section.left_margin.twips - section.right_margin.twips
        paragraph = formatted.paragraphs[-1]
        region = available - int(paragraph._p.pPr.ind.get(qn("w:left")))
        self.assertEqual(paragraph.text, date)
        self.assertGreater(region, 152 * 20)  # Previous output's under-sized region.
        self.assertGreaterEqual(region, (len(date) + 2) * 16 * 20)
        self.assertEqual(paragraph.runs[0].font.size.pt, 16)
        self.assertEqual(len(paragraph._p.findall('.//' + qn('w:br'))), 0)

    def test_signature_width_reserves_fullwidth_budget_and_proportional_headroom(self) -> None:
        available = 12000
        for text in ("2026年8月28日印发", "2026年12月28日印发", "２０２６年８月２８日印发", "WWW事务协调办公室"):
            with self.subTest(text=text):
                region = available - signature_indent_twips([text], 16, available)
                self.assertGreaterEqual(region, (len(text) + 2) * 16 * 20)
        region = available - signature_indent_twips(["某" * 20], 16, available)
        self.assertGreaterEqual(region, 23 * 16 * 20)
        self.assertEqual(signature_indent_twips(["某" * 50], 16, available), 0)

    def test_signature_width_ignores_combining_and_zero_width_marks(self) -> None:
        self.assertEqual(signature_indent_twips(["e\u0301单位"], 16, 10000), signature_indent_twips(["é单位"], 16, 10000))
        self.assertEqual(signature_indent_twips(["某\u200b单位"], 16, 10000), signature_indent_twips(["某单位"], 16, 10000))

    def test_signature_reformat_reuses_one_blank_and_preserves_suffix(self) -> None:
        for suffix in ("", "印发", "印发。"):
            for combined in (False, True):
                with self.subTest(suffix=suffix, combined=combined):
                    document = Document()
                    for text in ["示例标题", "正文。", "", ""]:
                        document.add_paragraph(text)
                    if combined:
                        document.add_paragraph(f"某单位办公室\n2026年8月28日{suffix}")
                    else:
                        document.add_paragraph("某单位办公室")
                        document.add_paragraph(f"2026年8月28日{suffix}")
                    output = self.root / "signature-repeat.docx"
                    for _ in range(2):
                        finalize(document, output, load_preset(), classify_paragraphs(document), [])
                        document = Document(output)
                        self.assertEqual([p.text for p in document.paragraphs], [
                            "示例标题", "", "正文。", "", "某单位办公室", f"2026年8月28日{suffix}",
                        ])

    def test_long_signature_expands_to_text_width_without_negative_indent(self) -> None:
        document = Document()
        classifications = add_text(document, "示例标题\n正文。\n[落款]" + "示例" * 25 + "办公室\n[落款]某年某月某日")
        output = self.root / "long-signature.docx"
        finalize(document, output, load_preset(), classifications, [])
        for p in Document(output).paragraphs[-3:]:
            self.assertEqual(p._p.pPr.ind.get(qn("w:left")), "0")

    def test_retired_colophon_map_preserves_non_signature_content(self) -> None:
        document = Document()
        document.add_paragraph("示例标题")
        document.add_paragraph("正文。")
        p = document.add_paragraph("抄送：某单位。")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].font.size = Pt(12)
        p.runs[0].font.color.rgb = RGBColor(31, 78, 121)
        before = p._p.xml
        classifications = classify_paragraphs(document)
        classifications[-1]["role"] = "colophon"
        output = self.root / "retired-map.docx"
        finalize(document, output, load_preset(), classifications, [])
        self.assertEqual(classifications[-1]["role"], "skip")
        self.assertEqual(Document(output).paragraphs[-1]._p.xml, before)

    def test_validator_rejects_fixed_signature_indent_and_duplicate_spacing(self) -> None:
        document = Document()
        classifications = add_text(document, "示例标题\n正文。\n某单位办公室 2026年8月28日印发")
        output = self.root / "signature-bad-geometry.docx"
        finalize(document, output, load_preset(), classifications, [])
        document = Document(output)
        document.paragraphs[-1]._p.pPr.ind.set(qn("w:leftChars"), "2600")
        spacer = document.paragraphs[-3]._p
        spacer.addprevious(deepcopy(spacer))
        document.save(output)
        scrub_docx(output)
        report = validate_document(output, load_preset())
        self.assertFalse(report["valid"])
        self.assertTrue(any("right-side block positioning" in error for error in report["errors"]))
        self.assertTrue(any("more than one preceding blank" in error for error in report["errors"]))

    def test_existing_document_preserves_table_and_picture(self) -> None:
        source = self.root / "source.docx"
        image = self.root / "pixel.png"
        image.write_bytes(PNG_1X1)
        document = Document()
        title = document.add_paragraph("示例标题")
        title.runs[0].font.color.rgb = RGBColor(31, 78, 121)
        document.add_paragraph("这是正文。")
        table = document.add_table(rows=2, cols=1)
        table.cell(0, 0).text = "表头"
        table.cell(1, 0).text = "表格内容"
        document.add_picture(str(image))
        document.save(source)
        inspection = inspect_document(source)
        self.assertEqual(inspection["summary"]["tables"], 1)
        self.assertGreaterEqual(inspection["summary"]["drawings"], 1)
        mapping = self.root / "classification.json"
        mapping.write_text(
            json.dumps({"classifications": inspection["classifications"], "tables": inspection["tables"]}, ensure_ascii=False),
            encoding="utf-8",
        )
        output = self.root / "formatted.docx"
        self.run_script("docx_engine.py", "format-existing", "--input", source, "--output", output, "--classification", mapping, "--confirmed")
        formatted = Document(output)
        self.assertEqual(formatted.tables[0].cell(1, 0).text, "表格内容")
        self.assertEqual(run_color(formatted.paragraphs[0].runs[0]), "000000")
        self.assertIn(run_east_asia_font(formatted.tables[0].cell(0, 0).paragraphs[0].runs[0]), {"方正黑体_GBK", "黑体", "SimHei"})
        self.assertIn(run_east_asia_font(formatted.tables[0].cell(1, 0).paragraphs[0].runs[0]), {"方正仿宋_GBK", "仿宋_GB2312", "FangSong_GB2312", "FangSong"})
        self.assertEqual(run_color(formatted.tables[0].cell(0, 0).paragraphs[0].runs[0]), "000000")
        self.assertEqual(run_color(formatted.tables[0].cell(1, 0).paragraphs[0].runs[0]), "000000")
        self.assertTrue(formatted.tables[0].cell(0, 0).paragraphs[0].runs[0].font.bold)
        self.assertTrue(formatted.tables[0].cell(1, 0).paragraphs[0].runs[0].font.bold)
        with zipfile.ZipFile(source) as before, zipfile.ZipFile(output) as after:
            before_media = {name: before.read(name) for name in before.namelist() if name.startswith("word/media/")}
            after_media = {name: after.read(name) for name in after.namelist() if name.startswith("word/media/")}
        self.assertEqual(before_media, after_media)

    def test_table_title_and_two_header_rows_use_role_fonts(self) -> None:
        source = self.root / "two-level-table.docx"
        document = Document()
        document.add_paragraph("示例公文标题")
        document.add_paragraph("这是表格前的正文。")
        table_title = document.add_paragraph("测试任务安排表")
        table_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = document.add_table(rows=4, cols=2)
        table.cell(0, 0).merge(table.cell(0, 1)).text = "第一层表头"
        table.cell(1, 0).text = "第二层甲"
        table.cell(1, 1).text = "第二层乙"
        table.cell(2, 0).text = "正文甲"
        table.cell(2, 1).text = "正文乙"
        table.cell(3, 0).text = "正文丙"
        table.cell(3, 1).text = "正文丁"
        document.save(source)

        inspection = inspect_document(source)
        self.assertEqual(inspection["tables"][0]["title_paragraph_index"], 2)
        self.assertEqual(inspection["tables"][0]["header_rows"], 2)
        mapping = self.root / "two-level-table-map.json"
        mapping.write_text(
            json.dumps({"classifications": inspection["classifications"], "tables": inspection["tables"]}, ensure_ascii=False),
            encoding="utf-8",
        )
        output = self.root / "two-level-table-formatted.docx"
        self.run_script("docx_engine.py", "format-existing", "--input", source, "--output", output, "--classification", mapping, "--confirmed")

        formatted = Document(output)
        title_run = next(paragraph.runs[0] for paragraph in formatted.paragraphs if paragraph.text == "测试任务安排表")
        header1_run = formatted.tables[0].cell(0, 0).paragraphs[0].runs[0]
        header2_run = formatted.tables[0].cell(1, 0).paragraphs[0].runs[0]
        body_run = formatted.tables[0].cell(2, 0).paragraphs[0].runs[0]
        self.assertIn(run_east_asia_font(title_run), {"方正小标宋_GBK", "方正小标宋简体"})
        self.assertIn(run_east_asia_font(header1_run), {"方正黑体_GBK", "黑体", "SimHei"})
        self.assertIn(run_east_asia_font(header2_run), {"方正楷体_GBK", "楷体_GB2312", "楷体GB2312", "KaiTi_GB2312", "KaiTi"})
        self.assertIn(run_east_asia_font(body_run), {"方正仿宋_GBK", "仿宋_GB2312", "仿宋GB2312", "FangSong_GB2312", "FangSong"})
        self.assertTrue(all(run_color(run) == "000000" for run in (title_run, header1_run, header2_run, body_run)))
        self.assertTrue(all(run.font.bold for run in (title_run, header1_run, header2_run, body_run)))

    def test_table_roles_follow_disabled_global_bold(self) -> None:
        source = self.root / "two-level-table-no-bold.docx"
        document = Document()
        document.add_paragraph("示例公文标题")
        document.add_paragraph("这是表格前的正文。")
        table_title = document.add_paragraph("测试任务安排表")
        table_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = document.add_table(rows=3, cols=2)
        table.cell(0, 0).merge(table.cell(0, 1)).text = "第一层表头"
        table.cell(1, 0).text = "第二层甲"
        table.cell(1, 1).text = "第二层乙"
        table.cell(2, 0).text = "正文甲"
        table.cell(2, 1).text = "正文乙"
        document.save(source)

        inspection = inspect_document(source)
        self.assertEqual(inspection["tables"][0]["title_paragraph_index"], 2)
        self.assertEqual(inspection["tables"][0]["header_rows"], 2)
        mapping = self.root / "two-level-table-no-bold-map.json"
        mapping.write_text(
            json.dumps({"classifications": inspection["classifications"], "tables": inspection["tables"]}, ensure_ascii=False),
            encoding="utf-8",
        )
        overrides = self.root / "no-bold-overrides.json"
        overrides.write_text(json.dumps({"global": {"bold": False}}), encoding="utf-8")
        output = self.root / "two-level-table-no-bold-formatted.docx"
        self.run_script(
            "docx_engine.py",
            "format-existing",
            "--input",
            source,
            "--output",
            output,
            "--classification",
            mapping,
            "--overrides",
            overrides,
            "--confirmed",
        )

        formatted = Document(output)
        runs = (
            next(paragraph.runs[0] for paragraph in formatted.paragraphs if paragraph.text == "测试任务安排表"),
            formatted.tables[0].cell(0, 0).paragraphs[0].runs[0],
            formatted.tables[0].cell(1, 0).paragraphs[0].runs[0],
            formatted.tables[0].cell(2, 0).paragraphs[0].runs[0],
        )
        self.assertTrue(all(run.font.bold is False for run in runs))

    def test_classification_cleanup_requires_confirmation_and_matching_hash(self) -> None:
        source = self.root / "cleanup-source.docx"
        output = self.root / "cleanup-output.docx"
        document = Document()
        document.add_paragraph("示例标题")
        document.save(source)
        document.save(output)
        mapping = self.root / "classification.json"
        payload = {
            "classifications": [{"index": 0, "role": "main_title"}],
            "tables": [{"index": 0, "header_rows": 1}],
        }
        mapping.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

        eligible = json.loads(
            self.run_script(
                "cleanup_classification.py",
                "--classification",
                mapping,
                "--input",
                source,
                "--output",
                output,
            ).stdout
        )
        self.assertTrue(eligible["eligible"])
        self.assertFalse(eligible["deleted"])
        self.assertTrue(eligible["confirmation_required"])
        self.assertTrue(mapping.exists())

        stale_hash = eligible["sha256"]
        mapping.write_text(json.dumps({**payload, "reviewed": True}, ensure_ascii=False), encoding="utf-8")
        refused = self.run_script(
            "cleanup_classification.py",
            "--classification",
            mapping,
            "--input",
            source,
            "--output",
            output,
            "--confirmed",
            "--expected-sha256",
            stale_hash,
            check=False,
        )
        self.assertNotEqual(refused.returncode, 0)
        self.assertTrue(mapping.exists())

        current = json.loads(
            self.run_script(
                "cleanup_classification.py",
                "--classification",
                mapping,
                "--input",
                source,
                "--output",
                output,
            ).stdout
        )
        deleted = json.loads(
            self.run_script(
                "cleanup_classification.py",
                "--classification",
                mapping,
                "--input",
                source,
                "--output",
                output,
                "--confirmed",
                "--expected-sha256",
                current["sha256"],
            ).stdout
        )
        self.assertTrue(deleted["deleted"])
        self.assertFalse(mapping.exists())

    def test_package_uses_declared_skill_name_as_archive_root(self) -> None:
        output = self.root / "renamed-clone.zip"
        result = json.loads(self.run_script("package_skill.py", "--output", output).stdout)
        self.assertTrue(result["created"])
        with zipfile.ZipFile(output) as archive:
            names = archive.namelist()
        self.assertTrue(names)
        self.assertTrue(all(name.startswith("official-document-formatting/") for name in names))
        self.assertIn("official-document-formatting/SKILL.md", names)

    def test_legacy_extension_and_tracked_changes_are_blocked(self) -> None:
        document = Document()
        document.add_paragraph("示例标题")
        source = self.root / "source.docx"
        document.save(source)
        legacy = self.root / "source.docm"
        legacy.write_bytes(source.read_bytes())
        self.assertTrue(inspect_document(legacy)["blockers"])
        tracked = self.root / "tracked.docx"
        with zipfile.ZipFile(source) as before, zipfile.ZipFile(tracked, "w") as after:
            for info in before.infolist():
                data = before.read(info.filename)
                if info.filename == "word/document.xml":
                    data = data.replace(b"<w:body>", b'<w:body><w:ins w:id="1"/>', 1)
                after.writestr(info, data)
        self.assertTrue(any("Tracked changes" in item for item in inspect_document(tracked)["blockers"]))

    def test_privacy_scrub_removes_authorship(self) -> None:
        output = self.root / "metadata.docx"
        document = Document()
        document.add_paragraph("示例标题")
        document.core_properties.author = "Sample Author"
        document.core_properties.last_modified_by = "Sample Editor"
        document.save(output)
        scrub_docx(output)
        with zipfile.ZipFile(output) as archive:
            core = archive.read("docProps/core.xml")
        self.assertNotIn(b"Sample Author", core)
        self.assertNotIn(b"Sample Editor", core)
        self.assertNotIn(b"lastModifiedBy", core)


if __name__ == "__main__":
    unittest.main(verbosity=2)
